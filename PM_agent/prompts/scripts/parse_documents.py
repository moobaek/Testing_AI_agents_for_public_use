"""
PM 문서 파싱 도구 (Document Parser)
=====================================
DOCX, XLSX, XLS 등 다양한 형식의 PM 문서를 파싱하여 JSON으로 변환

사용법:
    python parse_documents.py --input <폴더경로> --output <출력폴더>
    python parse_documents.py --file <파일경로> --type <문서유형>
"""

import os
import sys
import json
import argparse
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional


# ============================================================
# 파서 클래스들
# ============================================================

class DocxParser:
    """DOCX 파일 파서 (회의록, 과업지시서 등)"""
    
    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        try:
            from docx import Document
        except ImportError:
            return {"error": "python-docx 설치 필요: pip install python-docx"}
        
        doc = Document(file_path)
        
        # 문단 추출
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # 테이블 추출
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            "filename": os.path.basename(file_path),
            "type": "docx",
            "parsed_at": datetime.now().isoformat(),
            "content": {
                "paragraphs": paragraphs,
                "tables": tables,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables)
            }
        }


class XlsxParser:
    """XLSX/XLS 파일 파서 (견적서, 주간보고, 이슈리스트)"""
    
    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        try:
            import pandas as pd
        except ImportError:
            return {"error": "pandas 설치 필요: pip install pandas openpyxl xlrd"}
        
        # 엔진 선택
        if file_path.endswith('.xls'):
            engine = 'xlrd'
        else:
            engine = 'openpyxl'
        
        try:
            sheets = pd.read_excel(file_path, sheet_name=None, engine=engine)
        except Exception as e:
            return {"error": f"엑셀 파싱 실패: {str(e)}"}
        
        content = {}
        for sheet_name, df in sheets.items():
            # NaN 값 처리
            df = df.fillna('')
            
            # 헤더와 데이터 분리
            headers = df.columns.tolist()
            rows = df.values.tolist()
            
            content[sheet_name] = {
                "headers": headers,
                "rows": rows,
                "row_count": len(rows),
                "column_count": len(headers)
            }
        
        return {
            "filename": os.path.basename(file_path),
            "type": "xlsx" if file_path.endswith('.xlsx') else "xls",
            "parsed_at": datetime.now().isoformat(),
            "content": content,
            "sheet_count": len(sheets)
        }


class MdParser:
    """Markdown 파일 파서"""
    
    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 섹션 추출 (## 헤더 기준)
        sections = {}
        current_section = "intro"
        current_content = []
        
        for line in content.split('\n'):
            if line.startswith('## '):
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = line[3:].strip()
                current_content = []
            else:
                current_content.append(line)
        
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return {
            "filename": os.path.basename(file_path),
            "type": "md",
            "parsed_at": datetime.now().isoformat(),
            "content": {
                "raw": content,
                "sections": sections,
                "line_count": len(content.split('\n'))
            }
        }


# ============================================================
# 문서 유형별 파서 매핑
# ============================================================

DOCUMENT_PARSERS = {
    # 회의록
    "meeting": {
        "patterns": ["회의록", "Meeting", "미팅", "아젠다", "Agenda"],
        "parser": DocxParser,
        "template": "2_project_execution/04_meeting/meeting_template.md"
    },
    # 주간보고
    "weekly_report": {
        "patterns": ["주간", "Weekly", "WR-"],
        "parser": XlsxParser,
        "template": "2_project_execution/01_status_report/weekly_report.md"
    },
    # 견적서
    "quotation": {
        "patterns": ["견적", "Quotation", "Quote"],
        "parser": XlsxParser,
        "template": "1_project_initiation/05_quotation/quotation_template.md"
    },
    # 이슈리스트
    "issue": {
        "patterns": ["이슈", "Issue", "리스트"],
        "parser": XlsxParser,
        "template": "2_project_execution/05_issue_list/issue_list.md"
    },
    # 과업지시서
    "sow": {
        "patterns": ["과업", "SOW", "지시서"],
        "parser": DocxParser,
        "template": "1_project_initiation/07_sow/sow_template.md"
    },
    # 계약서
    "contract": {
        "patterns": ["계약", "Contract"],
        "parser": DocxParser,
        "template": "1_project_initiation/06_contract/contract_template.md"
    }
}


def detect_document_type(filename: str) -> Optional[str]:
    """파일명으로 문서 유형 감지"""
    for doc_type, config in DOCUMENT_PARSERS.items():
        for pattern in config["patterns"]:
            if pattern.lower() in filename.lower():
                return doc_type
    return None


def parse_file(file_path: str, doc_type: Optional[str] = None) -> Dict[str, Any]:
    """단일 파일 파싱"""
    filename = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    # 문서 유형 감지
    if not doc_type:
        doc_type = detect_document_type(filename)
    
    # 확장자별 파서 선택
    if ext in ['.docx']:
        result = DocxParser.parse(file_path)
    elif ext in ['.xlsx', '.xls']:
        result = XlsxParser.parse(file_path)
    elif ext in ['.md']:
        result = MdParser.parse(file_path)
    else:
        return {
            "filename": filename,
            "error": f"지원하지 않는 형식: {ext}",
            "supported": [".docx", ".xlsx", ".xls", ".md"]
        }
    
    # 문서 유형 및 템플릿 정보 추가
    result["document_type"] = doc_type
    if doc_type and doc_type in DOCUMENT_PARSERS:
        result["template"] = DOCUMENT_PARSERS[doc_type]["template"]
    
    return result


def parse_folder(folder_path: str, output_path: str) -> Dict[str, Any]:
    """폴더 내 모든 문서 파싱"""
    results = {
        "metadata": {
            "source_folder": folder_path,
            "parsed_at": datetime.now().isoformat(),
            "total_files": 0,
            "success": 0,
            "failed": 0
        },
        "documents": []
    }
    
    supported_extensions = ['.docx', '.xlsx', '.xls', '.md']
    
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            # 임시 파일 제외
            if file.startswith('~$'):
                continue
            
            ext = os.path.splitext(file)[1].lower()
            if ext in supported_extensions:
                file_path = os.path.join(root, file)
                result = parse_file(file_path)
                
                results["metadata"]["total_files"] += 1
                if "error" not in result:
                    results["metadata"]["success"] += 1
                else:
                    results["metadata"]["failed"] += 1
                
                results["documents"].append(result)
    
    # 결과 저장
    os.makedirs(output_path, exist_ok=True)
    output_file = os.path.join(output_path, "parsed_all_documents.json")
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    print(f"✅ 파싱 완료: {results['metadata']['success']}/{results['metadata']['total_files']}개")
    print(f"📁 출력: {output_file}")
    
    return results


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="PM 문서 파싱 도구")
    parser.add_argument("--input", "-i", help="입력 폴더 경로")
    parser.add_argument("--file", "-f", help="단일 파일 경로")
    parser.add_argument("--output", "-o", default="./temp", help="출력 폴더 (기본: ./temp)")
    parser.add_argument("--type", "-t", help="문서 유형 (meeting, weekly_report, quotation, issue, sow, contract)")
    
    args = parser.parse_args()
    
    if args.file:
        result = parse_file(args.file, args.type)
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif args.input:
        parse_folder(args.input, args.output)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
