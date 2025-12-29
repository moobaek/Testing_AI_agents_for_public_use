"""
회의록 전용 파서 (Meeting Parser)
=================================
DOCX 회의록에서 액션 아이템 및 주요 정보 추출

사용법:
    python parse_meeting.py --file <파일경로>
"""

import os
import sys
import json
import re
from datetime import datetime
from typing import Dict, List, Any, Optional


def parse_meeting_docx(file_path: str) -> Dict[str, Any]:
    """회의록 DOCX 파싱 및 구조화"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx 설치 필요: pip install python-docx"}
    
    doc = Document(file_path)
    filename = os.path.basename(file_path)
    
    # 날짜 추출 (파일명에서)
    date_match = re.search(r'(\d{6}|\d{8})', filename)
    meeting_date = None
    if date_match:
        date_str = date_match.group(1)
        if len(date_str) == 6:
            meeting_date = f"20{date_str[:2]}-{date_str[2:4]}-{date_str[4:6]}"
        else:
            meeting_date = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
    
    # 문단 분석
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # 테이블 분석
    tables = []
    action_items = []
    participants = []
    decisions = []
    
    for table in doc.tables:
        table_data = []
        headers = []
        
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            
            if i == 0:
                headers = row_data
            else:
                table_data.append(row_data)
                
                # 액션 아이템 감지
                row_text = ' '.join(row_data).lower()
                if any(kw in row_text for kw in ['해야', '진행', '완료', '담당', 'action', 'todo']):
                    action_items.append({
                        "content": row_data,
                        "headers": headers
                    })
        
        tables.append({
            "headers": headers,
            "rows": table_data
        })
    
    # 결정 사항 추출
    for p in paragraphs:
        p_lower = p.lower()
        if any(kw in p_lower for kw in ['결정', 'decision', '합의', '확정']):
            decisions.append(p)
    
    # 액션 아이템 유형 분류
    categorized_actions = {
        "요구사항변경": [],
        "일정변경": [],
        "이슈발생": [],
        "기타": []
    }
    
    for action in action_items:
        action_text = ' '.join(action["content"]).lower()
        if any(kw in action_text for kw in ['기능', '스펙', '요구', '설계', '변경']):
            categorized_actions["요구사항변경"].append(action)
        elif any(kw in action_text for kw in ['일정', '납기', '마일스톤', '지연']):
            categorized_actions["일정변경"].append(action)
        elif any(kw in action_text for kw in ['이슈', '버그', '문제', '오류']):
            categorized_actions["이슈발생"].append(action)
        else:
            categorized_actions["기타"].append(action)
    
    # 연동 제안 생성
    sync_suggestions = []
    if categorized_actions["요구사항변경"]:
        sync_suggestions.append({
            "type": "Document_Update_Checker",
            "reason": "요구사항 변경 감지",
            "items": len(categorized_actions["요구사항변경"])
        })
    if categorized_actions["일정변경"]:
        sync_suggestions.append({
            "type": "Progress_Tracker",
            "reason": "일정 변경 감지",
            "items": len(categorized_actions["일정변경"])
        })
    if categorized_actions["이슈발생"]:
        sync_suggestions.append({
            "type": "Troubleshooting_Management",
            "reason": "이슈 발생 감지",
            "items": len(categorized_actions["이슈발생"])
        })
    
    return {
        "filename": filename,
        "type": "meeting",
        "parsed_at": datetime.now().isoformat(),
        "meeting_date": meeting_date,
        "content": {
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables,
            "table_count": len(tables)
        },
        "extracted": {
            "decisions": decisions,
            "action_items": action_items,
            "action_item_count": len(action_items),
            "categorized_actions": categorized_actions
        },
        "sync_suggestions": sync_suggestions,
        "template": "2_project_execution/04_meeting/meeting_template.md"
    }


def main():
    import argparse
    parser = argparse.ArgumentParser(description="회의록 파서")
    parser.add_argument("--file", "-f", required=True, help="회의록 파일 경로")
    parser.add_argument("--output", "-o", help="출력 파일 경로")
    
    args = parser.parse_args()
    
    result = parse_meeting_docx(args.file)
    
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"✅ 저장: {args.output}")
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
