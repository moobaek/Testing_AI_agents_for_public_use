"""
회의록 파서
"""

import os
import re
from datetime import datetime
from typing import Dict, List, Any


def parse_meeting_docx(file_path: str) -> Dict[str, Any]:
    """회의록 DOCX 파싱 및 액션아이템 추출"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx 설치 필요"}
    
    doc = Document(file_path)
    filename = os.path.basename(file_path)
    
    # 날짜 추출
    date_match = re.search(r'(\d{6}|\d{8})', filename)
    meeting_date = None
    if date_match:
        d = date_match.group(1)
        if len(d) == 6:
            meeting_date = f"20{d[:2]}-{d[2:4]}-{d[4:6]}"
        else:
            meeting_date = f"{d[:4]}-{d[4:6]}-{d[6:8]}"
    
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    tables = []
    action_items = []
    
    for table in doc.tables:
        headers = []
        table_data = []
        
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            
            if i == 0:
                headers = row_data
            else:
                table_data.append(row_data)
                
                row_text = ' '.join(row_data).lower()
                if any(kw in row_text for kw in ['해야', '진행', '완료', '담당', 'action']):
                    action_items.append({"content": row_data, "headers": headers})
        
        tables.append({"headers": headers, "rows": table_data})
    
    # 액션아이템 분류
    categorized = {"요구사항변경": [], "일정변경": [], "이슈발생": [], "기타": []}
    
    for action in action_items:
        text = ' '.join(action["content"]).lower()
        if any(kw in text for kw in ['기능', '스펙', '요구', '설계']):
            categorized["요구사항변경"].append(action)
        elif any(kw in text for kw in ['일정', '납기', '마일스톤']):
            categorized["일정변경"].append(action)
        elif any(kw in text for kw in ['이슈', '버그', '문제']):
            categorized["이슈발생"].append(action)
        else:
            categorized["기타"].append(action)
    
    # 연동 제안
    sync = []
    if categorized["요구사항변경"]:
        sync.append({"type": "Document_Update_Checker", "reason": "요구사항 변경 감지"})
    if categorized["일정변경"]:
        sync.append({"type": "Progress_Tracker", "reason": "일정 변경 감지"})
    if categorized["이슈발생"]:
        sync.append({"type": "Troubleshooting_Management", "reason": "이슈 발생 감지"})
    
    return {
        "filename": filename,
        "type": "meeting",
        "parsed_at": datetime.now().isoformat(),
        "meeting_date": meeting_date,
        "content": {"paragraphs": paragraphs, "tables": tables},
        "extracted": {
            "action_items": action_items,
            "categorized_actions": categorized
        },
        "sync_suggestions": sync,
        "template": "2_project_execution/04_meeting/meeting_template.md"
    }
