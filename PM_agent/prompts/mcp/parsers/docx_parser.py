"""
DOCX 파서 (Word 문서)
"""

import os
from datetime import datetime
from typing import Dict, Any


def parse_docx(file_path: str) -> Dict[str, Any]:
    """DOCX 파일 파싱"""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx 설치 필요: pip install python-docx"}
    
    doc = Document(file_path)
    
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    
    return {
        "filename": os.path.basename(file_path),
        "type": "docx",
        "parsed_at": datetime.now().isoformat(),
        "content": {
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables,
            "table_count": len(tables)
        }
    }
