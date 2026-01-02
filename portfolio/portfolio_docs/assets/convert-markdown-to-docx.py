#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX 변환 스크립트
한글과컴퓨터에서 열 수 있는 DOCX 파일 생성
"""

import sys
import os
import re
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    print(f"필요한 라이브러리가 설치되지 않았습니다: {e}")
    print("다음 명령어로 설치하세요: pip install python-docx")
    sys.exit(1)

try:
    from playwright.sync_api import sync_playwright
except ImportError:
    print("경고: playwright가 설치되지 않았습니다. Mermaid 다이어그램을 이미지로 변환할 수 없습니다.")
    print("다음 명령어로 설치하세요: pip install playwright")
    print("그리고 브라우저를 설치하세요: playwright install chromium")
    sync_playwright = None


def set_korean_font(run, size=11):
    """한글 폰트 설정"""
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)


def create_mermaid_html(mermaid_code):
    """Mermaid 다이어그램을 렌더링할 HTML 생성"""
    html_template = """<!DOCTYPE html>
<html lang="ko">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
    <style>
        body {{
            margin: 0;
            padding: 20px;
            background: white;
            font-family: 'Malgun Gothic', sans-serif;
        }}
        .mermaid {{
            text-align: center;
            background: white;
        }}
    </style>
</head>
<body>
    <div class="mermaid">
{mermaid_code}
    </div>
    <script>
        mermaid.initialize({{
            startOnLoad: true,
            theme: 'default',
            themeVariables: {{
                fontSize: '14px',
                primaryColor: '#2563eb',
                primaryTextColor: '#fff',
                primaryBorderColor: '#1d4ed8',
                lineColor: '#64748b',
                secondaryColor: '#f8fafc',
                tertiaryColor: '#e2e8f0'
            }}
        }});
    </script>
</body>
</html>"""
    return html_template.format(mermaid_code=mermaid_code)


def mermaid_to_image(mermaid_code, output_path):
    """Mermaid 다이어그램을 이미지로 변환"""
    if sync_playwright is None:
        return False
    
    try:
        html_content = create_mermaid_html(mermaid_code)
        
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            
            # HTML 콘텐츠 로드
            page.set_content(html_content, wait_until='networkidle')
            
            # Mermaid 렌더링 대기
            page.wait_for_selector('.mermaid svg', timeout=10000)
            
            # 추가 대기 (렌더링 완료 보장)
            page.wait_for_timeout(2000)
            
            # SVG 요소 찾기
            svg_element = page.query_selector('.mermaid svg')
            if svg_element:
                # SVG 요소의 스크린샷 캡처
                svg_element.screenshot(path=output_path)
            else:
                # 전체 페이지 스크린샷 (대체 방법)
                page.screenshot(path=output_path, full_page=True)
            
            browser.close()
        
        return os.path.exists(output_path)
    except Exception as e:
        print(f"Mermaid 이미지 변환 오류: {e}")
        return False


def insert_image_to_docx(doc, image_path, max_width_inches=6.0):
    """DOCX에 이미지 삽입"""
    try:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        # 이미지 삽입
        run.add_picture(image_path, width=Inches(max_width_inches))
        
        return True
    except Exception as e:
        print(f"이미지 삽입 오류: {e}")
        return False


def parse_inline_formatting(text, paragraph):
    """인라인 포맷팅 파싱 (강조, 링크, 코드 등)"""
    if not text:
        return
    
    # 단계별로 처리: 링크 -> 볼드 -> 이탤릭 -> 인라인 코드
    # 1단계: 링크 처리 [텍스트](URL)
    parts = re.split(r'(\[.*?\]\(.*?\))', text)
    for part in parts:
        if part.startswith('[') and '](' in part:
            # 링크
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text = match.group(1)
                # 링크 텍스트 내부의 포맷팅 처리 (간단한 방법)
                # 링크 텍스트에서 마크다운 제거 후 기본 링크로 처리
                # (복잡한 중첩 포맷팅은 일단 제외)
                clean_link_text = re.sub(r'\*\*|\*|`', '', link_text)
                run = paragraph.add_run(clean_link_text)
                set_korean_font(run)
                # 링크 스타일 적용
                try:
                    run._element.set(qn('w:rStyle'), 'Hyperlink')
                except:
                    pass
        elif part:
            # 2단계: 볼드 처리 **텍스트**
            bold_parts = re.split(r'(\*\*.*?\*\*)', part)
            for bold_part in bold_parts:
                if bold_part.startswith('**') and bold_part.endswith('**') and len(bold_part) > 4:
                    # 볼드 내부 텍스트 (이탤릭, 코드 등 처리)
                    bold_text = bold_part[2:-2]
                    bold_sub_parts = re.split(r'(\*[^*].*?\*|`.*?`)', bold_text)
                    for bold_sub in bold_sub_parts:
                        if bold_sub.startswith('*') and bold_sub.endswith('*') and not bold_sub.startswith('**'):
                            # 이탤릭
                            run = paragraph.add_run(bold_sub[1:-1])
                            set_korean_font(run)
                            run.font.bold = True
                            run.font.italic = True
                        elif bold_sub.startswith('`') and bold_sub.endswith('`'):
                            # 인라인 코드 (볼드 적용)
                            run = paragraph.add_run(bold_sub[1:-1])
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(30, 41, 59)
                        elif bold_sub:
                            # 일반 볼드 텍스트
                            run = paragraph.add_run(bold_sub)
                            set_korean_font(run)
                            run.font.bold = True
                elif bold_part:
                    # 3단계: 이탤릭 처리 *텍스트* (볼드가 아닌 경우)
                    italic_parts = re.split(r'(\*[^*\n].*?\*)', bold_part)
                    for italic_part in italic_parts:
                        if italic_part.startswith('*') and italic_part.endswith('*') and not italic_part.startswith('**'):
                            # 이탤릭 내부 텍스트 (코드 등 처리)
                            italic_text = italic_part[1:-1]
                            code_match = re.search(r'(`.*?`)', italic_text)
                            if code_match:
                                # 이탤릭과 코드 혼합
                                code_start = italic_text.find('`')
                                if code_start > 0:
                                    run = paragraph.add_run(italic_text[:code_start])
                                    set_korean_font(run)
                                    run.font.italic = True
                                code_end = italic_text.find('`', code_start + 1)
                                if code_end > code_start:
                                    run = paragraph.add_run(italic_text[code_start+1:code_end])
                                    run.font.name = 'Consolas'
                                    run.font.size = Pt(9)
                                    run.font.italic = True
                                    run.font.color.rgb = RGBColor(30, 41, 59)
                                if code_end < len(italic_text) - 1:
                                    run = paragraph.add_run(italic_text[code_end+1:])
                                    set_korean_font(run)
                                    run.font.italic = True
                            else:
                                # 순수 이탤릭
                                run = paragraph.add_run(italic_text)
                                set_korean_font(run)
                                run.font.italic = True
                        elif italic_part:
                            # 4단계: 인라인 코드 처리 `코드`
                            code_parts = re.split(r'(`.*?`)', italic_part)
                            for code_part in code_parts:
                                if code_part.startswith('`') and code_part.endswith('`'):
                                    run = paragraph.add_run(code_part[1:-1])
                                    run.font.name = 'Consolas'
                                    run.font.size = Pt(9)
                                    run.font.color.rgb = RGBColor(30, 41, 59)
                                elif code_part:
                                    # 일반 텍스트
                                    run = paragraph.add_run(code_part)
                                    set_korean_font(run)


def markdown_to_docx(md_file, output_file):
    """Markdown 파일을 DOCX로 변환"""
    
    # Markdown 파일 읽기
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Document 생성
    doc = Document()
    
    # Markdown 파싱
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_rows = []
    in_code_block = False
    code_lang = None
    temp_image_files = []  # 임시 이미지 파일 목록
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 코드 블록 처리
        if stripped.startswith('```'):
            if in_code_block:
                # 코드 블록 종료 (일반 코드 블록만 여기서 처리, Mermaid는 시작 시 처리)
                if code_lang != 'mermaid':
                    # 일반 코드 블록 처리
                    code_text = '\n'.join(code_lines)
                    if code_text.strip():
                        p = doc.add_paragraph(code_text)
                        p.style = 'No Spacing'
                        for run in p.runs:
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(30, 41, 59)
                
                in_code_block = False
                code_lang = None
                code_lines = []
            else:
                # 코드 블록 시작
                code_lang = stripped[3:].strip()
                if code_lang == 'mermaid':
                    # Mermaid 다이어그램 코드 수집
                    in_code_block = True
                    code_lines = []
                    i += 1
                    # 닫는 ```를 찾을 때까지 코드 수집
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        code_lines.append(lines[i])
                        i += 1
                    # 닫는 ```를 만났으므로 여기서 처리
                    if i < len(lines):
                        # Mermaid 다이어그램을 이미지로 변환
                        mermaid_code = '\n'.join(code_lines)
                        if mermaid_code.strip():
                            # 임시 이미지 파일 생성
                            temp_image = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                            temp_image.close()
                            temp_image_path = temp_image.name
                            temp_image_files.append(temp_image_path)
                            
                            # Mermaid를 이미지로 변환
                            if mermaid_to_image(mermaid_code, temp_image_path):
                                # DOCX에 이미지 삽입
                                insert_image_to_docx(doc, temp_image_path)
                            else:
                                # 변환 실패 시 안내 메시지
                                p = doc.add_paragraph()
                                run = p.add_run("[Mermaid 다이어그램: 변환 실패]")
                                set_korean_font(run)
                                run.font.italic = True
                                run.font.color.rgb = RGBColor(128, 128, 128)
                        # 닫는 ``` 건너뛰기
                        i += 1
                    # 코드 블록 상태 초기화
                    in_code_block = False
                    code_lang = None
                    code_lines = []
                    continue
                else:
                    in_code_block = True
                    code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # 표 처리
        if '|' in stripped and stripped.count('|') >= 2:
            if not in_table:
                in_table = True
                table_rows = []
            
            # 표 헤더 구분선 건너뛰기
            if re.match(r'^\|[\s\-\|:]+\|$', stripped):
                i += 1
                continue
            
            # 표 행 파싱
            cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        else:
            # 표 종료
            if in_table and table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        p = cell.paragraphs[0]
                        parse_inline_formatting(cell_data, p)
                        for run in p.runs:
                            set_korean_font(run, 10)
                
                in_table = False
                table_rows = []
        
        # 제목 처리
        if stripped.startswith('# '):
            title_text = stripped[2:].strip()
            p = doc.add_heading('', level=1)
            parse_inline_formatting(title_text, p)
            for run in p.runs:
                set_korean_font(run, 20)
                if not run.font.bold:
                    run.font.bold = True
        elif stripped.startswith('## '):
            title_text = stripped[3:].strip()
            p = doc.add_heading('', level=2)
            parse_inline_formatting(title_text, p)
            for run in p.runs:
                set_korean_font(run, 16)
                if not run.font.bold:
                    run.font.bold = True
                run.font.color.rgb = RGBColor(37, 99, 235)
        elif stripped.startswith('### '):
            title_text = stripped[4:].strip()
            p = doc.add_heading('', level=3)
            parse_inline_formatting(title_text, p)
            for run in p.runs:
                set_korean_font(run, 13)
                if not run.font.bold:
                    run.font.bold = True
        elif stripped.startswith('#### '):
            title_text = stripped[5:].strip()
            p = doc.add_heading('', level=4)
            parse_inline_formatting(title_text, p)
            for run in p.runs:
                set_korean_font(run, 11)
                if not run.font.bold:
                    run.font.bold = True
        # 구분선
        elif stripped.startswith('---'):
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 목록
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(text, p)
            # 포맷팅이 적용되지 않은 경우 기본 텍스트 추가
            if not p.runs:
                run = p.add_run(text)
                set_korean_font(run)
            else:
                for run in p.runs:
                    if not hasattr(run.font, 'name') or not run.font.name:
                        set_korean_font(run)
        # 번호 목록
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped).strip()
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(text, p)
            # 포맷팅이 적용되지 않은 경우 기본 텍스트 추가
            if not p.runs:
                run = p.add_run(text)
                set_korean_font(run)
            else:
                for run in p.runs:
                    if not hasattr(run.font, 'name') or not run.font.name:
                        set_korean_font(run)
        # 일반 텍스트
        elif stripped:
            p = doc.add_paragraph()
            parse_inline_formatting(stripped, p)
            # 포맷팅이 적용되지 않은 경우 기본 텍스트 추가
            if not p.runs:
                run = p.add_run(stripped)
                set_korean_font(run)
            else:
                # 포맷팅이 적용된 경우에도 폰트 확인
                for run in p.runs:
                    if not hasattr(run.font, 'name') or not run.font.name:
                        set_korean_font(run)
        
        i += 1
    
    # 문서 저장
    doc.save(output_file)
    
    # 임시 이미지 파일 정리
    for temp_file in temp_image_files:
        try:
            if os.path.exists(temp_file):
                os.remove(temp_file)
        except Exception as e:
            print(f"임시 파일 삭제 오류 ({temp_file}): {e}")
    
    print(f"DOCX 파일 생성 완료: {output_file}")
    if temp_image_files:
        print(f"Mermaid 다이어그램 {len(temp_image_files)}개를 이미지로 변환하여 삽입했습니다.")
    print("\n다음 단계:")
    print("1. 한글과컴퓨터 한글 프로그램 실행")
    print(f"2. '{output_file}' 파일 열기")
    print("3. 파일 > 다른 이름으로 저장 > HWP 형식 선택")
    print(f"4. '{output_file.replace('.docx', '.hwp')}'로 저장")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("사용법: python convert-markdown-to-docx.py <input.md> [output.docx]")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else md_file.replace('.md', '.docx')
    
    if not os.path.exists(md_file):
        print(f"파일을 찾을 수 없습니다: {md_file}")
        sys.exit(1)
    
    markdown_to_docx(md_file, output_file)
